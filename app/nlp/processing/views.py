import os

import spacy
from django.contrib import messages
from django.db import IntegrityError
from django.http import HttpResponse
from django.shortcuts import render, redirect
from django.conf import settings
from django.views import View
from docx import Document
from spacy.matcher import Matcher

from .extractor import TextExtractor
from . import utils
from .forms import TextForm, UploadResumeModelForm
from .models import Resume


class Main(View):
    def get(self, request):
        input_form = TextForm()
        upload_form = UploadResumeModelForm()
        return render(request, 'index.html', {'input_form': input_form, 'upload_form': upload_form})

    def post(self, request):

        if request.method == "POST":
            input_form = TextForm(request.POST)
            Resume.objects.all().delete()
            upload_form = UploadResumeModelForm(request.POST, request.FILES)
            files = request.FILES.getlist('resume')
            resumes_data = []
            if input_form.is_valid():
                try:
                    # saving the file
                    resume = Resume(resume=files)
                    resume.save()

                    # extracting resume entities
                    parser = TextExtractor('', input_form['text'].value())
                    data = parser.get_extracted_data()
                    resumes_data.append(data)
                    resume.name = data.get('name')
                    resume.email = data.get('email')
                    resume.mobile_number = data.get('mobile_number')
                    resume.education = data.get('college_name')
                    # if data.get('degree') is not None:
                    #     resume.education = ', '.join(data.get('degree'))
                    # else:
                    #     resume.education = None
                    resume.company_names = data.get('company_names')
                    resume.college_name = data.get('education')
                    resume.designation = data.get('designation')
                    resume.total_experience = data.get('total_experience')
                    if data.get('skills') is not None:
                        resume.skills = ', '.join(data.get('skills'))
                    else:
                        resume.skills = None
                    if data.get('experience') is not None:
                        resume.experience = ', '.join(data.get('experience'))
                    else:
                        resume.experience = None
                    resume.save()
                except IntegrityError:
                    messages.warning(request, 'Такое имя уже существует:', file.name)
                        # return redirect('#')
                resumes = Resume.objects.all()
                messages.success(request, 'Успешно загружено')
                context = {
                    "resumes": resumes
                }
                return render(request, 'index.html', context)
                # resumes_data.append(input_form['text'].value())
                # # document.add_heading(utils.extract_name(self.__nlp, self.__matcher)
                # #                      + " \n"+ utils.extract_email(input_form['text'].value()))
                # document.add_heading(utils.extract_skills(self.__nlp, self.__noun_chunks))
                #
                # response = HttpResponse(
                #     content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                # response['Content-Disposition'] = 'attachment; filename=download.docx'
                # document.save(response)
                # messages.success(request, u"Успешно")
                # return response
                # # return render(request, 'index.html', {"form": form})

            elif upload_form.is_valid():
                for file in files:
                    try:
                        # saving the file
                        resume = Resume(resume=file)
                        resume.save()

                        # extracting resume entities
                        parser = TextExtractor(os.path.join(settings.MEDIA_ROOT, resume.resume.name))
                        data = parser.get_extracted_data()
                        resumes_data.append(data)
                        resume.name = data.get('name')
                        resume.email = data.get('email')
                        resume.mobile_number = data.get('mobile_number')
                        resume.education = data.get('college_name')
                        # if data.get('degree') is not None:
                        #     resume.education = ', '.join(data.get('degree'))
                        # else:
                        #     resume.education = None
                        resume.company_names = data.get('company_names')
                        resume.college_name = data.get('education')
                        resume.designation = data.get('designation')
                        resume.total_experience = data.get('total_experience')
                        if data.get('skills') is not None:
                            resume.skills = ', '.join(data.get('skills'))
                        else:
                            resume.skills = None
                        if data.get('experience') is not None:
                            resume.experience = ', '.join(data.get('experience'))
                        else:
                            resume.experience = None
                        resume.save()
                    except IntegrityError:
                        messages.warning(request, 'Такое имя уже существует:', file.name)
                        # return redirect('#')
                resumes = Resume.objects.all()
                messages.success(request, 'Успешно загружено')
                context = {
                    "resumes": resumes
                }
                return render(request, 'index.html', context)
        else:
            input_form = TextForm()
            upload_form = UploadResumeModelForm()
        return render(request, 'index.html', {"input_form": input_form, "upload_form": upload_form})



def download_docx(text):
    document = Document()
    document.add_heading(text, 0)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response
